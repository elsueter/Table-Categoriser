from docx import Document
from docx.shared import Inches

def write(document_name, document_title, contents):
    document = Document()

    document.add_heading(document_title, 0)

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    hdr_cells[3].text = 'Name'
    for qty, id, desc, name in contents:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc
        row_cells[3].text = name

    document.add_page_break()

    document.save(document_name)