from docx import Document


class TableStructure:
    name = ""
    contents = []

    def __init__(self, input_name):
        self.name = input_name
        self.contents = []

    def push(self, input_array):
        self.contents.append(input_array)


class Tables:
    categories = []

    def set_categories(self, input_categories):
        for category in input_categories:
            self.categories.append(TableStructure(category))

    def push(self, input_category, input_array):
        hit = False
        for category in self.categories:
            input_category = input_category.replace(" ", "")
            input_category = input_category.replace("-", "")
            input_category = input_category.lower()
            if category.name in input_category:
                category.push(input_array)
                hit = True
        if not hit:
            print(input_category)

    def print(self, input_category):
        print(f"Printing {input_category}:\n")
        for category in self.categories:
            if input_category == category.name:
                for row in category.contents:
                    print(row)



def read_contents(document_name):
    word_doc = Document(document_name)

    document_table = Tables()

    temp_codes = word_doc.tables[0].rows[0].cells[0].text
    temp_codes = temp_codes.replace(" ", "")
    temp_codes = temp_codes.replace("-", "")
    temp_codes = temp_codes.lower()

    codes = temp_codes.split("\n")

    codes.remove("codes")

    document_table.set_categories(codes)

    for table in word_doc.tables:
        for row in table.rows:
            temp_array = []
            for cell in row.cells:
                temp_array.append(cell.text)
            document_table.push(temp_array[0], temp_array)

    return document_table
