from docx import Document
from docx.shared import Inches

def write(document_name, document_title, contents):
    document = Document()

    document.add_heading(document_title, 0)

    #cols 

    table = document.add_table(rows=1, cols=len(contents[0]))
    hdr_cells = table.rows[0].cells
    for i in range(len(contents[0])):
        header = contents[0][i].split("\n")
        hdr_cells[i].text = {header[0]}

    for codes in contents[1:]:
        row_cells = table.add_row().cells
        for i in range(len(codes)):
            row_cells[i].text = str(codes[i])


    document.add_page_break()

    document.save(document_name)